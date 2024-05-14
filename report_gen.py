import docx

def create_rent_doc(data:list):
    """Создание выходного документа: счет"""
    doc = docx.Document()
    doc.add_heading('Счет за авто', 0)
    for row in data:
        doc.add_paragraph(f'Арендованная машина: {row[1]}')
        doc.add_paragraph(f'Клиент: {row[2]}')
        doc.add_paragraph(f'Дата начала аренды: {row[3].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Дата конца аренды: {row[4].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'К оплате: {row[5]}')
        doc.add_paragraph(f'Подпись клиента:_________________')
    doc.add_heading('', 0)
    doc.save('reports\\invoice.docx')


def create_service_doc(data:list):
    """Создание выходного документа: ТО"""
    doc = docx.Document()
    doc.add_heading('Техническое обслуживание', 0)
    for row in data:
        doc.add_paragraph(f'Машина: {row[1]}')
        doc.add_paragraph(f'Организация партнер: {row[2]}')
        doc.add_paragraph(f'Адресс: {row[3]}')
        doc.add_paragraph(f'Вид работы: {row[4]}')
        doc.add_paragraph(f'Дата последней работы: {row[5].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Подпись отвественного лица:_________________')
    doc.add_heading('', 0)
    doc.save('reports\\TO.docx')


def create_claim_doc(data:list):
    """Создание выходного документа: претензия"""
    doc = docx.Document()
    doc.add_heading('Претензия клиенту', 0)
    for row in data:
        doc.add_paragraph(f'Арендованная машина: {row[1]}')
        doc.add_paragraph(f'Клиент: {row[2]}')
        doc.add_paragraph(f'Дата жалобы: {row[3].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Причина жалобы: {row[4]}')
        doc.add_paragraph(f'Подпись клиента:_________________')
    doc.add_heading('', 0)
    doc.save('reports\\claim.docx')